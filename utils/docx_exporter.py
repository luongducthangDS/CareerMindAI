import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "667eea")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header(doc, title: str):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    _set_font(run, size=10, bold=True, color=(102, 126, 234))
    _add_horizontal_line(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def export_cv_docx(data: dict) -> bytes:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Header: Name ────────────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(data.get("ho_ten", "").upper())
    _set_font(name_run, size=22, bold=True, color=(31, 41, 55))
    name_p.paragraph_format.space_after = Pt(2)

    if data.get("chuc_danh"):
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(data["chuc_danh"])
        _set_font(title_run, size=12, color=(102, 126, 234))
        title_p.paragraph_format.space_after = Pt(4)

    # Contact line
    contacts = []
    for key, label in [("email", None), ("phone", None), ("linkedin", "LinkedIn"), ("github", "GitHub"), ("dia_chi", None)]:
        val = data.get(key, "").strip()
        if val:
            contacts.append(f"{label}: {val}" if label else val)

    if contacts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run("  |  ".join(contacts))
        _set_font(contact_run, size=9, color=(107, 114, 128))
        contact_p.paragraph_format.space_after = Pt(6)

    # ── Summary ─────────────────────────────────────────────────────────────────
    if data.get("tom_tat", "").strip():
        _section_header(doc, "Tóm tắt bản thân")
        p = doc.add_paragraph(data["tom_tat"].strip())
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            _set_font(run, size=10)

    # ── Experience ──────────────────────────────────────────────────────────────
    experiences = [e for e in data.get("kinh_nghiem", []) if e.get("cong_ty") or e.get("vi_tri")]
    if experiences:
        _section_header(doc, "Kinh nghiệm làm việc")
        for exp in experiences:
            # Company + Duration
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            run_company = p.add_run(exp.get("cong_ty", ""))
            _set_font(run_company, size=11, bold=True)
            if exp.get("thoi_gian"):
                run_time = p.add_run(f"  —  {exp['thoi_gian']}")
                _set_font(run_time, size=10, color=(107, 114, 128))

            # Position
            if exp.get("vi_tri"):
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                run_pos = p2.add_run(exp["vi_tri"])
                _set_font(run_pos, size=10, bold=True, color=(102, 126, 234))

            # Description bullets
            if exp.get("mo_ta", "").strip():
                for line in exp["mo_ta"].strip().split("\n"):
                    line = line.strip().lstrip("-•").strip()
                    if line:
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.paragraph_format.left_indent = Cm(0.5)
                        bp.paragraph_format.space_after = Pt(1)
                        run = bp.add_run(line)
                        _set_font(run, size=10)

    # ── Education ───────────────────────────────────────────────────────────────
    educations = [e for e in data.get("hoc_van", []) if e.get("truong")]
    if educations:
        _section_header(doc, "Học vấn")
        for edu in educations:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            run_school = p.add_run(edu.get("truong", ""))
            _set_font(run_school, size=11, bold=True)
            if edu.get("thoi_gian"):
                run_time = p.add_run(f"  —  {edu['thoi_gian']}")
                _set_font(run_time, size=10, color=(107, 114, 128))

            details = []
            if edu.get("nganh"):
                details.append(edu["nganh"])
            if edu.get("gpa"):
                details.append(f"GPA: {edu['gpa']}")
            if details:
                p2 = doc.add_paragraph("  |  ".join(details))
                p2.paragraph_format.space_after = Pt(2)
                for run in p2.runs:
                    _set_font(run, size=10, color=(75, 85, 99))

    # ── Projects ────────────────────────────────────────────────────────────────
    projects = [p for p in data.get("du_an", []) if p.get("ten")]
    if projects:
        _section_header(doc, "Dự án cá nhân")
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            run_name = p.add_run(proj.get("ten", ""))
            _set_font(run_name, size=11, bold=True)
            if proj.get("cong_nghe"):
                run_tech = p.add_run(f"  |  {proj['cong_nghe']}")
                _set_font(run_tech, size=10, color=(102, 126, 234))

            if proj.get("mo_ta", "").strip():
                for line in proj["mo_ta"].strip().split("\n"):
                    line = line.strip().lstrip("-•").strip()
                    if line:
                        bp = doc.add_paragraph(style="List Bullet")
                        bp.paragraph_format.left_indent = Cm(0.5)
                        bp.paragraph_format.space_after = Pt(1)
                        run = bp.add_run(line)
                        _set_font(run, size=10)

            if proj.get("link"):
                pl = doc.add_paragraph()
                pl.paragraph_format.space_after = Pt(2)
                run_link = pl.add_run(f"🔗 {proj['link']}")
                _set_font(run_link, size=9, color=(102, 126, 234))

    # ── Skills ──────────────────────────────────────────────────────────────────
    if data.get("ky_nang", "").strip():
        _section_header(doc, "Kỹ năng")
        p = doc.add_paragraph(data["ky_nang"].strip())
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            _set_font(run, size=10)

    # ── Certifications ──────────────────────────────────────────────────────────
    if data.get("chung_chi", "").strip():
        _section_header(doc, "Chứng chỉ & Giải thưởng")
        p = doc.add_paragraph(data["chung_chi"].strip())
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            _set_font(run, size=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
